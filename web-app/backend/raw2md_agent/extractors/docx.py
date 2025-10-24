"""
DOCX text extraction for Raw2MD Agent.

Uses python-docx to extract text, tables, and headers from Word documents.
"""

import logging
from typing import Optional, Dict, Any

from .base import BaseExtractor, ExtractionError, CorruptedFileError

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX extraction will not work.")


class DOCXExtractor(BaseExtractor):
    """
    DOCX text extractor using python-docx.
    
    Features:
    - Extract all paragraphs and headers
    - Extract table content
    - Preserve document structure
    - Handle nested elements
    """
    
    def __init__(self, path: str, config: Optional[Dict[str, Any]] = None):
        """
        Initialize DOCX extractor.
        
        Args:
            path: Path to DOCX file
            config: Optional configuration dictionary
        """
        super().__init__(path, config)
        
        if not DOCX_AVAILABLE:
            raise ExtractionError("python-docx library not available", str(self.path))
    
    def extract(self) -> str:
        """
        Extract text from DOCX file.
        
        Returns:
            Extracted text content
            
        Raises:
            ExtractionError: If extraction fails
            CorruptedFileError: If DOCX is corrupted
        """
        try:
            doc = Document(str(self.path))
            logger.debug(f"Opened DOCX document")
            
            extracted_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Check if it's a heading
                    if paragraph.style and paragraph.style.name and paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        extracted_content.append(f"\n{'#' * int(level)} {text}\n")
                    else:
                        extracted_content.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    extracted_content.append(f"\n{table_text}\n")
            
            # Extract headers and footers
            for section in doc.sections:
                header_text = self._extract_header_footer(section.header)
                footer_text = self._extract_header_footer(section.footer)
                
                if header_text:
                    extracted_content.insert(0, f"{header_text}\n")
                if footer_text:
                    extracted_content.append(f"\n{footer_text}")
            
            full_text = "\n".join(extracted_content)
            
            if not full_text.strip():
                raise ExtractionError("No text content found in DOCX", str(self.path))
            
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            if "corrupted" in str(e).lower() or "invalid" in str(e).lower():
                raise CorruptedFileError(f"DOCX file appears to be corrupted: {e}", str(self.path), e)
            raise ExtractionError(f"Failed to extract text from DOCX: {e}", str(self.path), e)
    
    def _extract_table(self, table) -> str:
        """
        Extract text from a table.
        
        Args:
            table: python-docx table object
            
        Returns:
            Formatted table text
        """
        try:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    cells.append(cell_text)
                rows.append(" | ".join(cells))
            
            if rows:
                return "\n".join(rows)
            return ""
            
        except Exception as e:
            logger.warning(f"Error extracting table: {e}")
            return ""
    
    def _extract_header_footer(self, header_footer) -> str:
        """
        Extract text from header or footer.
        
        Args:
            header_footer: Header or footer object
            
        Returns:
            Header/footer text
        """
        try:
            paragraphs = []
            for paragraph in header_footer.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            return "\n".join(paragraphs)
            
        except Exception as e:
            logger.warning(f"Error extracting header/footer: {e}")
            return ""
    
    def get_paragraph_count(self) -> int:
        """
        Get the number of paragraphs in the document.
        
        Returns:
            Number of paragraphs
        """
        try:
            doc = Document(str(self.path))
            count = len(doc.paragraphs)
            return count
        except Exception as e:
            logger.warning(f"Error getting paragraph count: {e}")
            return 0
    
    def get_table_count(self) -> int:
        """
        Get the number of tables in the document.
        
        Returns:
            Number of tables
        """
        try:
            doc = Document(str(self.path))
            count = len(doc.tables)
            return count
        except Exception as e:
            logger.warning(f"Error getting table count: {e}")
            return 0
